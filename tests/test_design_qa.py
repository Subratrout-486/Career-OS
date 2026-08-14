from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from career_os.design_qa import audit_resume_design


def _make_resume_files(tmp_path: Path) -> dict[str, str]:
    pdf_path = tmp_path / "resume.pdf"
    pdf = canvas.Canvas(str(pdf_path), pagesize=A4)
    pdf.setFont("Helvetica", 11)
    y = A4[1] - 50
    for heading in ("PROFESSIONAL SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION"):
        pdf.drawString(50, y, heading)
        y -= 18
        pdf.drawString(50, y, "Evidence-grounded career content.")
        y -= 26
    pdf.save()

    docx_path = tmp_path / "resume.docx"
    document = Document()
    for heading in ("PROFESSIONAL SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION"):
        document.add_paragraph(heading)
        document.add_paragraph("Evidence-grounded career content.")
    document.save(str(docx_path))
    return {"pdf": str(pdf_path), "docx": str(docx_path)}


def test_design_qa_accepts_clean_single_page_artifacts(tmp_path):
    result = audit_resume_design(_make_resume_files(tmp_path))
    assert result["passed"] is True
    assert result["metrics"]["pdf_pages"] == 1
    assert result["metrics"]["docx_tables"] == 0


def test_design_qa_fails_closed_when_artifacts_are_missing():
    result = audit_resume_design({})
    assert result["passed"] is False
    assert "final PDF is missing" in result["issues"]
    assert "final DOCX is missing" in result["issues"]
