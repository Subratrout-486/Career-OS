import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem


def xml_safe(text) -> str:
    value = str(text or "")
    return "".join(ch for ch in value if ch in "\t\n\r" or ord(ch) >= 32)


def safe_filename(text: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "-", text.strip())
    return value.strip("-")[:90] or "career-os-resume"


def _experience_items(resume: dict):
    for item in resume.get("experience", []):
        if not isinstance(item, dict):
            yield str(item), "", []
            continue
        header = " — ".join(
            str(x).strip()
            for x in [item.get("title"), item.get("company"), item.get("dates")]
            if str(x or "").strip()
        )
        yield header or "Experience", "", item.get("bullets") or item.get("responsibilities") or []


def build_docx(resume: dict, output_path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run(xml_safe(resume.get("title") or "Resume"))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(xml_safe(text).upper())
        r.bold = True
        r.font.size = Pt(10.5)
        return p

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run(xml_safe(text))
        return p

    heading("Professional Summary")
    doc.add_paragraph(xml_safe(resume.get("summary") or ""))

    heading("Skills")
    doc.add_paragraph(xml_safe(" • ".join(str(x) for x in resume.get("skills", []) if str(x).strip())))

    heading("Experience")
    for header, _, bullets in _experience_items(resume):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(xml_safe(header))
        r.bold = True
        for item in bullets:
            bullet(item)

    heading("Education")
    for item in resume.get("education", []):
        bullet(item)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_pdf(resume: dict, output_path: Path) -> Path:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=15,
        leading=17, spaceAfter=5, alignment=1
    )
    heading_style = ParagraphStyle(
        "ResumeHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=9.5,
        leading=11, spaceBefore=5, spaceAfter=2
    )
    body_style = ParagraphStyle(
        "ResumeBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.7,
        leading=10.5, spaceAfter=2
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet", parent=body_style, leftIndent=9, firstLineIndent=-5, spaceAfter=1
    )

    story = [Paragraph(xml_safe(resume.get("title") or "Resume"), title_style)]

    def heading(text):
        story.append(Paragraph(text.upper(), heading_style))

    heading("Professional Summary")
    story.append(Paragraph(xml_safe(resume.get("summary") or ""), body_style))

    heading("Skills")
    story.append(Paragraph(xml_safe(" • ".join(str(x) for x in resume.get("skills", []) if str(x).strip())), body_style))

    heading("Experience")
    for header, _, bullets in _experience_items(resume):
        story.append(Paragraph(f"<b>{xml_safe(header)}</b>", body_style))
        for item in bullets:
            story.append(Paragraph("• " + xml_safe(item), bullet_style))

    heading("Education")
    for item in resume.get("education", []):
        story.append(Paragraph("• " + xml_safe(item), bullet_style))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4, rightMargin=14 * mm, leftMargin=14 * mm,
        topMargin=12 * mm, bottomMargin=12 * mm,
        title=xml_safe(resume.get("title") or "Resume")
    )
    doc.build(story)
    return output_path


def generate_resume_files(job: dict, resume: dict, output_dir: str = "generated_resumes") -> dict[str, str]:
    directory = Path(output_dir)
    company = safe_filename(str(job.get("company") or "Company"))
    role = safe_filename(str(job.get("title") or "Role"))
    stem = f"Subrat_Rout_{company}_{role}_Resume"
    docx_path = directory / f"{stem}.docx"
    pdf_path = directory / f"{stem}.pdf"
    build_docx(resume, docx_path)
    build_pdf(resume, pdf_path)
    return {"docx": str(docx_path), "pdf": str(pdf_path)}
